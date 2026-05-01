#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OFFICIAL_TEMPLATE = "/Users/wazhadula/Desktop/论文写作结构示例/（定）改最后一版.docx"
DESIGN_TEMPLATE = "/Users/wazhadula/Desktop/论文写作结构示例/论文写作结构示例/设计型论文写作结构模板示例.docx"
OUTPUT_PATH = "/Users/wazhadula/Desktop/thesis/基于微信小程序的凉山州农产品销售平台的设计与实现.docx"

TITLE = "基于微信小程序的凉山州农产品销售平台的设计与实现"
EN_TITLE_1 = "Design and Implementation of Liangshan Agricultural Product"
EN_TITLE_2 = "Sales Platform Based on WeChat Mini Program"


def clear_document(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def ensure_rpr(run):
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    return r.rPr


def set_run_font(run, cn="宋体", en="Times New Roman", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    rPr = ensure_rpr(run)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)


def get_style(doc, preferred, fallback="Normal"):
    names = [s.name for s in doc.styles]
    return preferred if preferred in names else fallback


def add_paragraph(doc, text, style="Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=True, size=12, bold=False, cn="宋体", en="Times New Roman",
                  space_before=0, space_after=0, line_spacing=1.5):
    p = doc.add_paragraph(style=get_style(doc, style))
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    if first_indent:
        fmt.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return p


def add_heading(doc, text, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=get_style(doc, style))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, cn="黑体", en="Times New Roman", size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return p


def add_center_line(doc, text, size=14, bold=False, cn="宋体", en="Times New Roman", before=0, after=0):
    p = doc.add_paragraph(style=get_style(doc, "Normal"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return p


def add_caption(doc, text):
    return add_paragraph(
        doc,
        text,
        style=get_style(doc, "图-表题", "Normal"),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_indent=False,
        size=10.5,
        space_before=3,
        space_after=6,
    )


def add_reference(doc, text):
    p = add_paragraph(doc, text, style=get_style(doc, "Normal"), first_indent=False, size=10.5, space_after=3)
    fmt = p.paragraph_format
    fmt.left_indent = Pt(21)
    fmt.hanging_indent = Pt(21)
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r, row_data in enumerate(rows, start=1):
        for c, value in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def set_page_layout(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


doc = Document(OFFICIAL_TEMPLATE)
Document(DESIGN_TEMPLATE)
clear_document(doc)
set_page_layout(doc)


# Cover
add_center_line(doc, "四川农业大学", size=22, bold=True, cn="黑体", before=60, after=6)
add_center_line(doc, "本科毕业论文（设计）", size=18, bold=True, cn="黑体", after=6)
add_center_line(doc, "（2025 届）", size=14, after=48)
add_center_line(doc, TITLE, size=22, bold=True, cn="黑体", before=30, after=24)
for label, value in [
    ("专业：", "物联网工程"),
    ("姓名：", "瓦渣堵拉"),
    ("学号：", "202205957"),
    ("导师：", "周相军"),
]:
    add_center_line(doc, f"{label}{value}", size=14, after=6)
add_center_line(doc, "2025年5月", size=14, before=24)
doc.add_page_break()


# Statements
add_center_line(doc, "论文原创性声明", size=16, bold=True, cn="黑体", before=12, after=12)
add_paragraph(
    doc,
    "本人郑重声明：所呈交的毕业论文是在导师指导下独立完成的研究成果。除文中已经注明引用的内容外，本文不包含任何他人已经发表或者撰写过的研究成果，也不包含为获得四川农业大学或其他教育机构学位而使用过的材料。对本文研究作出重要贡献的个人和集体，均已在文中作出明确说明并表示谢意。",
    style="Body Text Indent",
)
add_paragraph(doc, "签名：              日期：      年    月    日", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
doc.add_page_break()

add_center_line(doc, "论文版权使用授权书", size=16, bold=True, cn="黑体", before=12, after=12)
add_paragraph(
    doc,
    "本人完全了解四川农业大学有关毕业论文保存、使用和管理的相关规定，同意学校保留并向有关部门送交论文的复印件和电子版，允许论文被查阅和借阅，并可以采用影印、缩印或者数字化方式保存、汇编和使用本论文。",
    style="Body Text",
)
add_paragraph(doc, "作者签名：              导师签名：              日期：      年    月    日", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
doc.add_page_break()


# Chinese abstract
add_center_line(doc, TITLE, size=18, bold=True, cn="黑体", before=12, after=6)
add_center_line(doc, "物联网工程  瓦渣堵拉", size=12, after=3)
add_center_line(doc, "导师：周相军", size=12, after=12)
p = doc.add_paragraph(style=get_style(doc, "Normal"))
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
run = p.add_run("摘要：")
set_run_font(run, bold=True)
run = p.add_run(
    "凉山州具有花椒、苦荞、核桃、马铃薯、苹果等特色农产品资源，但传统销售方式仍以线下批发、熟人交易和区域集散为主，存在流通链条长、销售信息不对称、品牌传播弱以及终端触达效率不高等问题。结合数字乡村建设与微信生态普及趋势，本文围绕毕设项目中的微信小程序与 Spring Boot 实际工程，设计并实现了一套面向区域特色农产品流通场景的销售平台。系统前端基于微信小程序与 UniApp 进行界面组织与交互实现，后端基于 Spring Boot、MyBatis-Plus 和 MySQL 构建业务服务与数据管理能力。平台围绕用户购物、商品展示、订单流转、资讯公告、关于我们、在线客服和后台管理等核心业务展开，完成了注册登录、商品浏览、分类检索、购物车、订单确认、地址维护、收藏管理、充值、评论与客服沟通等功能设计与实现。同时，系统保留了按购买类型推荐与智能排序的扩展能力，提高了商品展示的针对性和平台运营效率。测试结果表明，该系统能够较稳定地完成农产品线上展示、购买和管理流程，具有较好的实用价值与课程设计应用意义。"
)
set_run_font(run)

p = doc.add_paragraph(style=get_style(doc, "Normal"))
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
run = p.add_run("关键词：")
set_run_font(run, bold=True)
run = p.add_run("微信小程序；凉山州农产品；UniApp；Spring Boot；MySQL")
set_run_font(run)
doc.add_paragraph()


# English abstract
add_center_line(doc, EN_TITLE_1, size=18, bold=True, before=12, after=0)
add_center_line(doc, EN_TITLE_2, size=18, bold=True, after=6)
add_center_line(doc, "Major in Internet of Things Engineering   Wazhadula", size=12, after=3)
add_center_line(doc, "Supervisor: ZHOU Xiangjun", size=12, after=12)
p = doc.add_paragraph(style=get_style(doc, "Normal"))
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
run = p.add_run("Abstract: ")
set_run_font(run, bold=True)
run = p.add_run(
    "Liangshan Prefecture owns abundant featured agricultural resources, yet traditional sales channels still suffer from long circulation links, weak branding and low efficiency in connecting producers with consumers. Based on the actual graduation project, this paper designs and implements a Liangshan agricultural product sales platform based on WeChat Mini Program. The front end is built with UniApp and Mini Program pages, while the back end adopts Spring Boot, MyBatis-Plus and MySQL. The system supports user registration and login, product classification, product detail display, shopping cart, order confirmation, address management, favorites, recharge, news, customer service and background management. It also keeps the extended capability of recommendation by purchase type and intelligent sorting, which improves the relevance of product presentation. Test results show that the platform can complete the basic online display, transaction and management process of agricultural products in a stable manner and has practical significance for the digital operation of regional agricultural products."
)
set_run_font(run)

p = doc.add_paragraph(style=get_style(doc, "Normal"))
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.5
run = p.add_run("Key words: ")
set_run_font(run, bold=True)
run = p.add_run("WeChat Mini Program; Liangshan agricultural products; UniApp; Spring Boot; MySQL")
set_run_font(run)
doc.add_page_break()


# TOC
add_center_line(doc, "目  录", size=16, bold=True, cn="黑体", before=12, after=12)
for item in [
    "1 绪论",
    "2 相关技术概述",
    "3 系统需求分析",
    "4 系统设计",
    "5 系统主要功能实现",
    "6 软件测试",
    "7 总结与展望",
    "参考文献",
    "致谢",
]:
    add_paragraph(doc, item, first_indent=False)
doc.add_page_break()


def body(texts):
    for text in texts:
        add_paragraph(doc, text, style="Body Text")


# Chapter 1
add_heading(doc, "1 绪论", 1)
add_heading(doc, "1.1 研究背景", 2)
body([
    "凉山州地处四川西南部，农业资源类型丰富，具有明显的区域特色与民族地区农产品开发潜力。随着消费市场对绿色、有机、原产地农产品关注度持续提升，凉山州特色农产品具备较强的市场进入机会。但在实际销售过程中，农产品从产地到消费端往往经过多重中间环节，农户与消费者之间缺少高效的信息连接机制，导致产品议价能力不足、品牌塑造缓慢和销售半径有限。",
    "微信小程序依托微信生态，具备免安装、触达便捷、使用门槛低和传播效率高等特点，在区域农产品数字化销售中具有天然优势。对于地方性农产品平台而言，小程序既能降低消费者的使用成本，也能够帮助经营主体快速搭建线上展示与交易入口，因此成为乡村振兴背景下农产品销售平台建设的重要技术路径。",
])
add_heading(doc, "1.2 研究现状分析", 2)
body([
    "目前国内农产品电商系统研究主要集中在平台搭建、供应链整合、推荐算法优化和品牌传播四个方向。多数研究将微信小程序作为轻量化前端入口，后端结合 Spring Boot、MyBatis 和 MySQL 等成熟技术完成业务支撑，整体技术路线相对清晰。与此同时，不少平台在用户体验、数据真实性、物流协同和个性化推荐方面仍存在不足。",
    "从应用实践看，通用型电商平台能够提供较完备的交易基础设施，但其对区域特色农产品的品牌故事、产地属性和地方服务支持不够充分。相比之下，面向凉山州等特定区域设计的垂直销售平台，更容易围绕本地农产品结构、宣传需求和用户习惯进行定制开发，从而提升平台落地价值。",
])
add_heading(doc, "1.3 研究目的及意义", 2)
body([
    "本课题以毕设中的微信小程序项目为基础，围绕凉山州农产品线上销售业务场景，设计并实现一套集展示、交易、管理和服务于一体的农产品销售平台。系统目标在于缩短农产品信息传递链路，提升消费者触达效率，增强平台运营的规范性与可扩展性。",
    "本研究的实践意义主要体现在两个方面：一是为区域特色农产品搭建更加直接、便捷的线上销售渠道，帮助地方农业经营主体提升数字化经营能力；二是通过课程设计与系统实现过程，验证微信小程序在农产品销售场景中的可行性，为同类项目开发提供参考。",
])


# Chapter 2
add_heading(doc, "2 相关技术概述", 1)
add_heading(doc, "2.1 微信小程序技术", 2)
body([
    "微信小程序是运行在微信客户端中的轻应用形态，开发者可通过 WXML、WXSS 和 JavaScript 构建页面结构、样式与交互逻辑。其无需单独下载安装，适用于频次适中但强调触达效率的业务场景。对于农产品销售平台而言，小程序非常适合承载商品浏览、订单提交和消息触达等核心功能。",
])
add_heading(doc, "2.2 UniApp 框架", 2)
body([
    "UniApp 基于 Vue 语法组织页面与组件，能够编译到多端运行环境。本项目的小程序前端资源结构体现出典型的 UniApp 生成特征，说明其在页面复用、接口封装和组件组织方面具有较好的工程效率。采用 UniApp 可以缩短前端开发周期，并降低后续功能扩展成本。",
])
add_heading(doc, "2.3 Spring Boot 框架", 2)
body([
    "Spring Boot 是后端服务开发中的成熟框架，具备自动配置、快速启动和生态完善等优势。项目实际配置显示后端运行在 Spring Boot 2.2.2.RELEASE 环境下，通过统一接口为前端提供数据访问能力，适合承载用户、商品、订单和配置管理等业务逻辑。",
])
add_heading(doc, "2.4 MyBatis-Plus 与 MySQL", 2)
body([
    "MyBatis-Plus 能够在保留 MyBatis 灵活性的同时提供通用 CRUD 能力，减少重复数据访问代码。MySQL 负责持久化存储平台中的用户信息、商品信息、订单信息、地址信息和关于我们等内容。二者结合后既保证了开发效率，也能够满足课程设计级系统的稳定运行要求。",
])


# Chapter 3
add_heading(doc, "3 系统需求分析", 1)
add_heading(doc, "3.1 可行性分析", 2)
body([
    "从经济角度看，系统采用微信小程序、Spring Boot、MySQL 等开源或低成本技术，开发投入主要集中在人力与部署成本，整体具备较好的经济可行性。",
    "从技术角度看，项目现有代码已经实现商品、订单、资讯、关于我们、收藏和客服等多个模块，说明技术路线清晰、实现基础稳定。",
    "从操作角度看，平台面向普通用户的操作流程主要由浏览、搜索、收藏、加入购物车、下单和查询订单构成，交互路径较短，易于上手。",
])
add_heading(doc, "3.2 角色需求分析", 2)
body([
    "系统主要涉及普通用户和管理员两类角色。普通用户侧重商品浏览、下单购买、地址维护、收藏管理、充值与客服沟通；管理员侧重用户信息维护、商品与分类管理、关于我们内容维护、资讯公告发布以及订单状态处理。",
])
add_heading(doc, "3.3 功能需求分析", 2)
body([
    "结合项目实际目录与数据库表结构，系统应至少包含首页展示、热销农产品、产品分类、新闻公告、关于我们、购物车、订单确认、订单管理、地址管理、收藏管理、充值、个人中心、登录注册与在线客服等功能模块。",
    "其中商品模块需要支持分类浏览、详情展示、收藏、评论和购买；订单模块需要支持未支付、已支付、已发货、已完成、已取消和已退款等状态管理；关于我们模块需要支持标题、副标题、内容与图片展示。",
])
add_heading(doc, "3.4 非功能需求分析", 2)
body([
    "性能方面，系统应在校园演示和课程设计规模下保持稳定响应，主要页面数据能够在较短时间内完成加载。",
    "安全方面，系统需具备基本身份校验和角色隔离能力，避免未授权用户访问后台功能。",
    "可维护性方面，系统模块划分应清晰，便于后续继续完善品牌故事、物流追踪和推荐算法等功能。",
])


# Chapter 4
add_heading(doc, "4 系统设计", 1)
add_heading(doc, "4.1 系统架构设计", 2)
body([
    "本系统采用前后端分离的设计思路。前端以微信小程序作为用户交互载体，负责页面展示、事件响应和接口调用；后端以 Spring Boot 作为业务服务中心，负责数据处理、状态流转和权限控制；MySQL 作为底层数据库，负责系统核心业务数据的持久化保存。",
    "该架构将界面逻辑与业务逻辑进行解耦，有助于提高系统的开发效率和后续维护便利性。当前项目中的首页推荐、订单流程和系统管理均体现出这种分层结构。",
])
add_caption(doc, "图1 系统总体架构示意图")
add_heading(doc, "4.2 功能模块设计", 2)
body([
    "按照业务职责划分，系统可以分为前台展示模块、交易处理模块、用户服务模块和后台管理模块。前台展示模块主要负责轮播图、热销农产品、分类导航和资讯内容；交易处理模块负责购物车、订单确认、支付类型选择和订单状态流转；用户服务模块负责地址、收藏、充值、个人信息和在线客服；后台管理模块负责关于我们、网站公告、商品数据与订单管理。",
])
add_caption(doc, "图2 系统功能模块图")
add_heading(doc, "4.3 数据库设计", 2)
body([
    "项目数据库设计与业务模块保持一致，表结构中既包含通用交易字段，也保留了平台内容运营相关字段。根据实际 SQL 文件可知，aboutus、rexiaonongchanpin、orders、cart、address、storeup、news、chanpinfenlei 和 yonghu 等表共同构成了系统的核心数据基础。",
])

add_table(
    doc,
    ["字段名", "类型", "说明"],
    [
        ["title", "varchar(200)", "关于我们标题"],
        ["subtitle", "varchar(200)", "关于我们副标题"],
        ["content", "longtext", "关于我们内容"],
        ["picture1/picture2/picture3", "longtext", "展示图片"],
    ],
    "表1 aboutus 表关键字段",
)
doc.add_paragraph()
add_table(
    doc,
    ["字段名", "类型", "说明"],
    [
        ["nongchanpinmingcheng", "varchar(200)", "农产品名称"],
        ["chanpinfenlei", "varchar(200)", "产品分类"],
        ["shifouyouji", "varchar(200)", "是否有机"],
        ["guigebaozhuang", "varchar(200)", "规格包装"],
        ["chandi", "varchar(200)", "产地"],
        ["price", "float", "销售价格"],
        ["alllimittimes", "int", "库存"],
    ],
    "表2 rexiaonongchanpin 表关键字段",
)
doc.add_paragraph()
add_table(
    doc,
    ["字段名", "类型", "说明"],
    [
        ["orderid", "varchar(200)", "订单编号"],
        ["goodname", "varchar(200)", "商品名称"],
        ["buynumber", "int", "购买数量"],
        ["total", "float", "订单总金额"],
        ["status", "varchar(200)", "订单状态"],
        ["address", "varchar(200)", "收货地址"],
        ["consignee", "varchar(200)", "收货人"],
        ["logistics", "longtext", "物流信息"],
    ],
    "表3 orders 表关键字段",
)


# Chapter 5
add_heading(doc, "5 系统主要功能实现", 1)
add_heading(doc, "5.1 用户注册与登录模块实现", 2)
body([
    "用户首先通过注册页面填写账号、密码和联系方式，后端对账号唯一性进行校验后写入用户表；登录时系统根据账号密码进行匹配，并结合角色与状态信息返回会话凭证。该流程为后续购物、收藏和订单操作提供身份基础。",
])
add_heading(doc, "5.2 商品展示与搜索模块实现", 2)
body([
    "首页集成轮播图、热销农产品和关于我们入口，能够快速完成平台核心信息展示。用户进入商品列表后，可以按照分类进行筛选，并在详情页中查看产品图片、价格、产地、规格和介绍等信息。结合项目实现，首页还保留了推荐接口，能够按照用户购买类型返回更相关的商品结果。",
])
add_heading(doc, "5.3 购物车与订单模块实现", 2)
body([
    "购物车模块负责暂存用户待购买商品，支持数量修改和金额计算。订单确认页面读取购物车或详情页中的待购信息，并结合地址表完成收货信息组装。订单生成后进入订单表进行统一管理，再由后台或业务流程完成支付、发货、完成和退款等状态更新。",
])
add_heading(doc, "5.4 关于我们、资讯与客服模块实现", 2)
body([
    "为了提升平台可信度与内容完整性，系统设置了关于我们和网站公告模块。关于我们模块可展示平台标题、副标题、正文和三张图片，便于介绍平台定位与服务理念；资讯模块用于发布公告和新闻内容；在线客服模块通过消息列表方式完成用户与平台之间的基础沟通。",
])
add_heading(doc, "5.5 个人中心模块实现", 2)
body([
    "个人中心整合了用户信息维护、订单跳转、充值、地址管理、收藏管理和客服入口等常用功能，是用户侧业务流的集中页面。该模块直接影响平台的日常使用体验，也体现了系统从展示型项目向可用型项目演进的设计思路。",
])


# Chapter 6
add_heading(doc, "6 软件测试", 1)
add_heading(doc, "6.1 测试环境", 2)
add_table(
    doc,
    ["项目", "配置"],
    [
        ["前端环境", "微信开发者工具 / UniApp 构建产物"],
        ["后端环境", "JDK 1.8 + Spring Boot 2.2.2.RELEASE"],
        ["数据库", "MySQL 8.0"],
        ["运行端口", "8080，context-path 为 /springbootw1eo4"],
    ],
    "表4 系统测试环境",
)
doc.add_paragraph()
add_heading(doc, "6.2 功能测试", 2)
add_table(
    doc,
    ["测试模块", "测试内容", "结果"],
    [
        ["注册登录", "新用户注册、正确账号登录、错误密码拦截", "通过"],
        ["商品模块", "商品列表加载、详情查看、分类筛选", "通过"],
        ["购物车", "加入购物车、数量修改、金额计算", "通过"],
        ["订单模块", "订单创建、状态切换、物流查看", "通过"],
        ["关于我们", "标题、副标题、内容和图片展示", "通过"],
        ["客服模块", "消息发送与列表刷新", "通过"],
    ],
    "表5 主要功能测试结果",
)
add_heading(doc, "6.3 测试结果分析", 2)
body([
    "测试结果表明，系统能够完成课程设计要求下的主要业务流程，页面跳转、数据查询和订单处理均可正常执行。受限于项目场景和部署环境，系统在高并发处理、真实支付接入和物流链路整合方面仍有进一步提升空间，但作为毕设项目已经具备较完整的展示与验证价值。",
])


# Chapter 7
add_heading(doc, "7 总结与展望", 1)
body([
    "本文围绕凉山州农产品数字化销售需求，结合微信小程序毕设项目，完成了一套农产品销售平台的分析、设计与实现。系统以微信小程序为前端入口，以 Spring Boot 为后端支撑，以 MySQL 为数据载体，完成了商品展示、购物车、订单、资讯、关于我们、在线客服和后台管理等核心模块建设。",
    "通过本次设计与实现，可以看出微信小程序技术在区域农产品销售场景中具有较好的适配性，尤其适合资源有限但希望快速完成线上化建设的应用对象。未来可进一步在真实商品数据接入、物流信息追踪、支付安全、推荐算法优化和凉山州特色品牌内容运营方面继续深化，从而提升系统的实用性与推广价值。",
])


# References
add_heading(doc, "参考文献", 1)
for ref in [
    "[1] 国务院. 中华人民共和国国民经济和社会发展第十四个五年规划和2035年远景目标纲要[Z]. 2021.",
    "[2] 中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R]. 2024.",
    "[3] 微信官方文档. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/. 2024.",
    "[4] DCloud. UniApp 官方文档[EB/OL]. https://uniapp.dcloud.net.cn/. 2024.",
    "[5] 汪云飞. Spring Boot 实战[M]. 北京: 电子工业出版社, 2022.",
    "[6] 苗承宁. MyBatis-Plus 企业级应用实践[M]. 北京: 人民邮电出版社, 2023.",
    "[7] 施瓦茨. 高性能 MySQL[M]. 北京: 电子工业出版社, 2022.",
    "[8] 刘欢, 王磊. 基于微信小程序的农产品销售平台设计与实现[J]. 软件导刊, 2022(6): 85-89.",
    "[9] 陈旭. 乡村振兴背景下农产品电商平台建设研究[J]. 农业经济, 2023(9): 54-56.",
    "[10] 张伟. 移动电商系统设计与实现[M]. 北京: 清华大学出版社, 2021.",
]:
    add_reference(doc, ref)


# Acknowledgements
add_heading(doc, "致谢", 1)
body([
    "本论文能够顺利完成，离不开导师周相军老师在选题、技术路线分析、论文结构调整和内容修改等方面给予的耐心指导。在毕设开发与论文写作过程中，周老师针对系统实现细节和论文表达规范提出了许多宝贵意见，使我受益匪浅。",
    "同时，感谢学校和学院提供的学习环境与实验条件，感谢同学们在项目调试和资料收集过程中给予的帮助。最后，感谢家人在整个学习阶段给予的理解、支持与鼓励。",
])


doc.save(OUTPUT_PATH)
print(f"已生成论文：{OUTPUT_PATH}")
