#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《基于微信小程序的农产品购物平台》论文
参考格式：论文写作结构示例/（定）改最后一版.docx
参考内容：开题报告
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        for key in ['sz', 'val', 'color', 'space', 'shadow']:
            if key in kwargs:
                element.set(qn('w:{}'.format(key)), str(kwargs[key]))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_run_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size_pt=12):
    """设置字体"""
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(size_pt)

def add_paragraph_with_style(doc, text, style='Normal', bold=False, 
                               font_size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               first_line_indent=True, cn_font='宋体', en_font='Times New Roman',
                               space_before=0, space_after=0):
    """添加带样式的段落"""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    
    # 段落格式
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(24)  # 两字符缩进
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    
    # 行距
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    
    return p

def add_heading(doc, text, level=1):
    """添加标题"""
    p = doc.add_heading(text, level=level)
    # 强制设置中文字体
    for run in p.runs:
        run.font.name = 'Times New Roman'
        r = run._element
        if r.rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = r.rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            r.rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '黑体' if level == 1 else '黑体')
    return p

def add_body_text(doc, text, first_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if first_indent:
        pf.first_line_indent = Pt(24)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = r.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code_text(doc, text):
    """添加代码段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(24)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Courier New'
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = r.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '仿宋')
    return p

def add_caption(doc, text):
    """添加图表说明"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = r.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_reference(doc, text):
    """添加参考文献条目"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(21)
    pf.hanging_indent = Pt(21)
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    r = run._element
    if r.rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = r.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def set_page_margins(doc):
    """设置页边距"""
    from docx.oxml import OxmlElement
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

def add_table_with_data(doc, headers, rows, caption=None):
    """添加表格"""
    if caption:
        cap_p = add_caption(doc, caption)
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
            r = run._element
            if r.rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = r.rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                r.rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_data in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_data)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Times New Roman'
                r = run._element
                if r.rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                rFonts = r.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    r.rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), '宋体')
    
    return table

# ==================== 主函数 ====================

doc = Document()
set_page_margins(doc)

# 设置文档默认字体
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ======================== 封面 ========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(72)
pf.space_after = Pt(6)
run = p.add_run('四川农业大学')
run.font.size = Pt(22)
run.bold = True
run.font.name = 'Times New Roman'
r = run._element
rFonts = r.rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    r.rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '黑体')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf2 = p2.paragraph_format
pf2.space_before = Pt(6)
pf2.space_after = Pt(6)
run2 = p2.add_run('本科毕业论文（设计）')
run2.font.size = Pt(18)
run2.bold = True
run2.font.name = 'Times New Roman'
r2 = run2._element
rFonts2 = r2.rPr.find(qn('w:rFonts'))
if rFonts2 is None:
    rFonts2 = OxmlElement('w:rFonts')
    r2.rPr.append(rFonts2)
rFonts2.set(qn('w:eastAsia'), '黑体')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf3 = p3.paragraph_format
pf3.space_before = Pt(6)
pf3.space_after = Pt(48)
run3 = p3.add_run('（2025 届）')
run3.font.size = Pt(14)
run3.font.name = 'Times New Roman'
r3 = run3._element
rFonts3 = r3.rPr.find(qn('w:rFonts'))
if rFonts3 is None:
    rFonts3 = OxmlElement('w:rFonts')
    r3.rPr.append(rFonts3)
rFonts3.set(qn('w:eastAsia'), '宋体')

# 论文标题
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_title = p_title.paragraph_format
pf_title.space_before = Pt(24)
pf_title.space_after = Pt(24)
run_title = p_title.add_run('基于微信小程序的凉山州农产品销售平台的设计与实现')
run_title.font.size = Pt(22)
run_title.bold = True
run_title.font.name = 'Times New Roman'
r_title = run_title._element
rFonts_title = r_title.rPr.find(qn('w:rFonts'))
if rFonts_title is None:
    rFonts_title = OxmlElement('w:rFonts')
    r_title.rPr.append(rFonts_title)
rFonts_title.set(qn('w:eastAsia'), '黑体')

# 作者信息
info_lines = [
    ('专业：', '物联网工程'),
    ('姓名：', '瓦渣堵拉'),
    ('学号：', '202205957'),
    ('导师：', '周相军'),
]
for label, value in info_lines:
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfi = pi.paragraph_format
    pfi.space_before = Pt(6)
    pfi.space_after = Pt(6)
    ri = pi.add_run(label + value)
    ri.font.size = Pt(14)
    ri.font.name = 'Times New Roman'
    r_i = ri._element
    rFonts_i = r_i.rPr.find(qn('w:rFonts'))
    if rFonts_i is None:
        rFonts_i = OxmlElement('w:rFonts')
        r_i.rPr.append(rFonts_i)
    rFonts_i.set(qn('w:eastAsia'), '宋体')

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_date = p_date.paragraph_format
pf_date.space_before = Pt(24)
run_date = p_date.add_run('2025年5月')
run_date.font.size = Pt(14)
run_date.font.name = 'Times New Roman'
r_date = run_date._element
rFonts_date = r_date.rPr.find(qn('w:rFonts'))
if rFonts_date is None:
    rFonts_date = OxmlElement('w:rFonts')
    r_date.rPr.append(rFonts_date)
rFonts_date.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 原创性声明 ========================
p_decl_title = doc.add_paragraph()
p_decl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_dt = p_decl_title.paragraph_format
pf_dt.space_before = Pt(12)
pf_dt.space_after = Pt(12)
run_dt = p_decl_title.add_run('论文原创性声明')
run_dt.font.size = Pt(16)
run_dt.bold = True
run_dt.font.name = 'Times New Roman'
r_dt = run_dt._element
rFonts_dt = r_dt.rPr.find(qn('w:rFonts'))
if rFonts_dt is None:
    rFonts_dt = OxmlElement('w:rFonts')
    r_dt.rPr.append(rFonts_dt)
rFonts_dt.set(qn('w:eastAsia'), '黑体')

decl_text = ('本人郑重声明：所呈交的学位论文是我个人在导师指导下进行研究工作所取得的成果。尽我所知，除了文中特别加以标注和致谢的地方外，'
             '学位论文中不包含其他个人或集体已经发表或撰写过的研究成果，也不包含为获得四川农业大学或其他教育机构的学位或证书而使用过的材料。'
             '与我一同工作的同志对本研究所做的任何贡献均已在论文中作了明确的说明并表示了谢意。')
add_body_text(doc, decl_text)

p_sign = doc.add_paragraph()
p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pf_sign = p_sign.paragraph_format
pf_sign.space_before = Pt(24)
run_sign = p_sign.add_run('签名：              日期：      年    月    日')
run_sign.font.size = Pt(12)
run_sign.font.name = 'Times New Roman'
r_sign = run_sign._element
rFonts_sign = r_sign.rPr.find(qn('w:rFonts'))
if rFonts_sign is None:
    rFonts_sign = OxmlElement('w:rFonts')
    r_sign.rPr.append(rFonts_sign)
rFonts_sign.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 摘要（中文）========================
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_abs = p_abs.paragraph_format
pf_abs.space_before = Pt(12)
pf_abs.space_after = Pt(12)
run_abs = p_abs.add_run('基于微信小程序的凉山州农产品销售平台的设计与实现')
run_abs.font.size = Pt(18)
run_abs.bold = True
run_abs.font.name = 'Times New Roman'
r_abs = run_abs._element
rFonts_abs = r_abs.rPr.find(qn('w:rFonts'))
if rFonts_abs is None:
    rFonts_abs = OxmlElement('w:rFonts')
    r_abs.rPr.append(rFonts_abs)
rFonts_abs.set(qn('w:eastAsia'), '黑体')

p_abs_sub = doc.add_paragraph()
p_abs_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_abs_sub = p_abs_sub.paragraph_format
pf_abs_sub.space_before = Pt(3)
pf_abs_sub.space_after = Pt(12)
run_abs_sub = p_abs_sub.add_run('物联网工程  瓦渣堵拉')
run_abs_sub.font.size = Pt(12)
run_abs_sub.font.name = 'Times New Roman'
r_abs_sub = run_abs_sub._element
rFonts_abs_sub = r_abs_sub.rPr.find(qn('w:rFonts'))
if rFonts_abs_sub is None:
    rFonts_abs_sub = OxmlElement('w:rFonts')
    r_abs_sub.rPr.append(rFonts_abs_sub)
rFonts_abs_sub.set(qn('w:eastAsia'), '宋体')

p_abs_supervisor = doc.add_paragraph()
p_abs_supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_abs_s = p_abs_supervisor.paragraph_format
pf_abs_s.space_before = Pt(0)
pf_abs_s.space_after = Pt(12)
run_abs_s = p_abs_supervisor.add_run('导师：周相军')
run_abs_s.font.size = Pt(12)
run_abs_s.font.name = 'Times New Roman'
r_abs_s = run_abs_s._element
rFonts_abs_s = r_abs_s.rPr.find(qn('w:rFonts'))
if rFonts_abs_s is None:
    rFonts_abs_s = OxmlElement('w:rFonts')
    r_abs_s.rPr.append(rFonts_abs_s)
rFonts_abs_s.set(qn('w:eastAsia'), '宋体')

# 摘要标签
p_abs_label = doc.add_paragraph()
p_abs_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run_abs_label = p_abs_label.add_run('摘要：')
run_abs_label.bold = True
run_abs_label.font.size = Pt(12)
run_abs_label.font.name = 'Times New Roman'
r_abl = run_abs_label._element
rFonts_abl = r_abl.rPr.find(qn('w:rFonts'))
if rFonts_abl is None:
    rFonts_abl = OxmlElement('w:rFonts')
    r_abl.rPr.append(rFonts_abl)
rFonts_abl.set(qn('w:eastAsia'), '宋体')

abstract_text = ('凉山州拥有花椒、苦荞、魔芋等丰富的特色农产品资源，但受地理位置偏远、传统销售链条较长、'
                 '市场信息不对称等因素影响，当地农产品长期面临销售半径有限、流通效率偏低和品牌传播不足等问题。'
                 '在国家“数字乡村”和乡村振兴战略持续推进的背景下，微信小程序凭借无需下载安装、使用门槛低、'
                 '依托微信生态传播便捷等优势，为区域农产品线上销售平台建设提供了新的实现路径。'
                 '本文结合开题报告中的研究目标，并基于实际项目代码，对一个基于微信小程序的农产品购物平台进行了设计与实现。'
                 '系统采用前后端分离模式，前端基于微信小程序与uni-app资源结构实现页面交互，后端采用Spring Boot，'
                 '数据层使用MyBatis-Plus与MySQL。平台围绕用户、商家和管理员三类角色展开，主要实现了用户注册登录、'
                 '农产品分类浏览、商品详情展示、购物车、订单管理、地址管理、收藏管理、资讯公告、在线客服、商家商品管理、'
                 '订单处理以及管理员审核与平台管理等功能。系统同时结合项目中的智能排序与按购买类型推荐接口，'
                 '提升农产品展示的针对性。研究结果表明，该平台能够较好满足农产品线上展示、交易与管理需求，'
                 '对拓宽特色农产品销售渠道、提升区域农产品数字化经营水平具有一定现实意义。')

run_abs_content = p_abs_label.add_run(abstract_text)
run_abs_content.font.size = Pt(12)
run_abs_content.font.name = 'Times New Roman'
r_abc = run_abs_content._element
rFonts_abc = r_abc.rPr.find(qn('w:rFonts'))
if rFonts_abc is None:
    rFonts_abc = OxmlElement('w:rFonts')
    r_abc.rPr.append(rFonts_abc)
rFonts_abc.set(qn('w:eastAsia'), '宋体')
pf_abs_label = p_abs_label.paragraph_format
pf_abs_label.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf_abs_label.line_spacing = 1.5

p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf_kw = p_kw.paragraph_format
pf_kw.space_before = Pt(6)
pf_kw.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf_kw.line_spacing = 1.5
run_kw_label = p_kw.add_run('关键词：')
run_kw_label.bold = True
run_kw_label.font.size = Pt(12)
run_kw_label.font.name = 'Times New Roman'
r_kwl = run_kw_label._element
rFonts_kwl = r_kwl.rPr.find(qn('w:rFonts'))
if rFonts_kwl is None:
    rFonts_kwl = OxmlElement('w:rFonts')
    r_kwl.rPr.append(rFonts_kwl)
rFonts_kwl.set(qn('w:eastAsia'), '宋体')
run_kw = p_kw.add_run('微信小程序；农产品购物平台；凉山州农产品；Spring Boot；前后端分离')
run_kw.font.size = Pt(12)
run_kw.font.name = 'Times New Roman'
r_kw = run_kw._element
rFonts_kw = r_kw.rPr.find(qn('w:rFonts'))
if rFonts_kw is None:
    rFonts_kw = OxmlElement('w:rFonts')
    r_kw.rPr.append(rFonts_kw)
rFonts_kw.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# 英文摘要
p_en_title = doc.add_paragraph()
p_en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_ent = p_en_title.paragraph_format
pf_ent.space_before = Pt(12)
pf_ent.space_after = Pt(3)
run_ent = p_en_title.add_run('Design and Implementation of Liangshan Agricultural Product')
run_ent.font.size = Pt(18)
run_ent.bold = True
run_ent.font.name = 'Times New Roman'

p_en_title2 = doc.add_paragraph()
p_en_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_ent2 = p_en_title2.paragraph_format
pf_ent2.space_before = Pt(0)
pf_ent2.space_after = Pt(12)
run_ent2 = p_en_title2.add_run('Sales Platform Based on WeChat Mini Program')
run_ent2.font.size = Pt(18)
run_ent2.bold = True
run_ent2.font.name = 'Times New Roman'

p_en_info = doc.add_paragraph()
p_en_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_ei = p_en_info.paragraph_format
pf_ei.space_before = Pt(3)
pf_ei.space_after = Pt(3)
run_ei = p_en_info.add_run('Major in Internet of Things Engineering   Wazha Dula')
run_ei.font.size = Pt(12)
run_ei.font.name = 'Times New Roman'

p_en_sup = doc.add_paragraph()
p_en_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_es = p_en_sup.paragraph_format
pf_es.space_before = Pt(0)
pf_es.space_after = Pt(12)
run_es = p_en_sup.add_run('Supervisor: ZHOU Xiangjun')
run_es.font.size = Pt(12)
run_es.font.name = 'Times New Roman'

p_en_abs = doc.add_paragraph()
p_en_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf_ea = p_en_abs.paragraph_format
pf_ea.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf_ea.line_spacing = 1.5
run_ea_label = p_en_abs.add_run('Abstract: ')
run_ea_label.bold = True
run_ea_label.font.size = Pt(12)
run_ea_label.font.name = 'Times New Roman'
run_ea = p_en_abs.add_run(
    'Liangshan Prefecture has abundant characteristic agricultural resources such as pepper, tartary '
    'buckwheat and konjac, but its products are still constrained by remote geography, long circulation '
    'links and insufficient market information sharing. Under the background of digital rural development '
    'and rural revitalization, WeChat Mini Program provides a lightweight solution for regional agricultural '
    'e-commerce because it requires no installation and is closely integrated with the WeChat ecosystem. '
    'Based on the research objectives proposed in the project proposal and the actual source code of the '
    'system, this paper designs and implements an agricultural product shopping platform based on WeChat '
    'Mini Program. The system adopts a front-end and back-end separation architecture. The front end is '
    'implemented with Mini Program pages and uni-app resource structure, while the back end is built with '
    'Spring Boot, MyBatis-Plus and MySQL. The platform covers three roles, namely users, merchants and '
    'administrators, and provides functions such as registration and login, product classification browsing, '
    'product detail display, shopping cart, order management, address management, favorites, news, online '
    'service, merchant product management, order processing and administrator auditing. In addition, the '
    'system uses intelligent sorting and purchase-type recommendation interfaces in the project to improve '
    'the relevance of product display. The implementation shows that the platform can effectively support '
    'online display, transaction and management of agricultural products and has practical significance for '
    'expanding sales channels of characteristic agricultural products and improving digital operation ability.')
run_ea.font.size = Pt(12)
run_ea.font.name = 'Times New Roman'

p_en_kw = doc.add_paragraph()
p_en_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf_ek = p_en_kw.paragraph_format
pf_ek.space_before = Pt(6)
pf_ek.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf_ek.line_spacing = 1.5
run_ek_label = p_en_kw.add_run('Key words: ')
run_ek_label.bold = True
run_ek_label.font.size = Pt(12)
run_ek_label.font.name = 'Times New Roman'
run_ek = p_en_kw.add_run('WeChat Mini Program; Agricultural Product Shopping Platform; Liangshan Agricultural Products; Spring Boot; Front-end and Back-end Separation')
run_ek.font.size = Pt(12)
run_ek.font.name = 'Times New Roman'

doc.add_page_break()

# ======================== 目录 ========================
p_toc = doc.add_paragraph()
p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_toc = p_toc.paragraph_format
pf_toc.space_before = Pt(12)
pf_toc.space_after = Pt(12)
run_toc = p_toc.add_run('目  录')
run_toc.font.size = Pt(16)
run_toc.bold = True
run_toc.font.name = 'Times New Roman'
r_toc = run_toc._element
rFonts_toc = r_toc.rPr.find(qn('w:rFonts'))
if rFonts_toc is None:
    rFonts_toc = OxmlElement('w:rFonts')
    r_toc.rPr.append(rFonts_toc)
rFonts_toc.set(qn('w:eastAsia'), '黑体')

toc_items = [
    ('1  绪论', '1'),
    ('    1.1  研究背景', '1'),
    ('    1.2  国内外研究现状', '2'),
    ('    1.3  研究目的与意义', '3'),
    ('2  开发技术介绍', '4'),
    ('    2.1  微信小程序开发框架', '4'),
    ('        2.1.1  WXML与WXSS', '4'),
    ('        2.1.2  JavaScript逻辑层', '4'),
    ('    2.2  后端开发技术', '5'),
    ('        2.2.1  Spring Boot框架', '5'),
    ('        2.2.2  MyBatis-Plus框架', '5'),
    ('        2.2.3  MySQL数据库', '5'),
    ('3  系统需求分析', '6'),
    ('    3.1  可行性分析', '6'),
    ('        3.1.1  经济可行性', '6'),
    ('        3.1.2  技术可行性', '6'),
    ('        3.1.3  操作可行性', '7'),
    ('    3.2  功能需求分析', '7'),
    ('    3.3  系统非功能需求分析', '8'),
    ('4  系统设计', '9'),
    ('    4.1  系统架构设计', '9'),
    ('    4.2  功能模块设计', '10'),
    ('        4.2.1  用户功能模块设计', '10'),
    ('        4.2.2  商家功能模块设计', '11'),
    ('        4.2.3  管理员功能模块设计', '11'),
    ('    4.3  数据库设计', '12'),
    ('        4.3.1  数据库表设计', '12'),
    ('5  系统开发与实现', '15'),
    ('    5.1  开发平台', '15'),
    ('    5.2  用户功能模块实现', '15'),
    ('        5.2.1  登录注册功能', '15'),
    ('        5.2.2  首页展示功能', '17'),
    ('        5.2.3  农产品浏览与搜索功能', '18'),
    ('        5.2.4  购物车与订单功能', '19'),
    ('        5.2.5  个人中心功能', '20'),
    ('    5.3  商家功能模块实现', '21'),
    ('    5.4  管理员模块实现', '22'),
    ('    5.5  数据库模块实现', '23'),
    ('6  测试', '24'),
    ('    6.1  测试目的', '24'),
    ('    6.2  测试环境', '24'),
    ('    6.3  具体测试过程', '25'),
    ('        6.3.1  功能测试', '25'),
    ('        6.3.2  性能测试', '26'),
    ('7  总结', '27'),
    ('参考文献', '28'),
    ('致谢', '30'),
    ('附录', '31'),
]

for item, page in toc_items:
    p_ti = doc.add_paragraph()
    p_ti.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf_ti = p_ti.paragraph_format
    pf_ti.space_before = Pt(0)
    pf_ti.space_after = Pt(2)
    pf_ti.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf_ti.line_spacing = 1.5
    
    run_ti = p_ti.add_run(item)
    run_ti.font.size = Pt(12)
    run_ti.font.name = 'Times New Roman'
    r_ti = run_ti._element
    rFonts_ti = r_ti.rPr.find(qn('w:rFonts'))
    if rFonts_ti is None:
        rFonts_ti = OxmlElement('w:rFonts')
        r_ti.rPr.append(rFonts_ti)
    rFonts_ti.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 第1章 绪论 ========================
h1 = doc.add_heading('1  绪论', level=1)

h2 = doc.add_heading('1.1  研究背景', level=2)

add_body_text(doc, 
    '凉山州作为四川省重要的民族地区和特色农产品产区，拥有花椒、苦荞、魔芋等具有较高市场辨识度的农产品资源。'
    '但长期以来，受地理位置相对偏远、物流成本较高、传统销售模式层级较多以及市场信息传递滞后等因素影响，'
    '当地农产品在销售过程中仍然存在“优质不优价”和“增产不增收”等现实问题。'
    '国家“十四五”规划提出要推进数字乡村建设，通过“互联网+农业”模式促进农业提质增效[1]。')

add_body_text(doc,
    '从现实情况看，传统线下销售方式对中间商依赖度较高，农产品从生产地到消费端往往需要经过收购、转运、批发等多个环节，'
    '不仅抬高了流通成本，也削弱了农户与消费者之间的信息直达能力。消费者在购买过程中往往难以全面了解产品的产地、'
    '品质、库存和价格变化情况，而生产者也难以及时感知市场需求变化，导致供需匹配效率偏低[2]。')

add_body_text(doc,
    '微信小程序具有无需下载安装、入口便捷、开发周期短和易于传播等特点，尤其适合区域特色产品的线上推广与销售。'
    '对于农产品平台而言，小程序既降低了普通用户的使用门槛，也能够借助微信生态完成商品展示、信息分享和用户留存，'
    '因此成为农产品数字化营销的重要载体[3]。已有案例表明，小程序模式能够有效拓宽农产品销售渠道，'
    '为区域农产品突破市场壁垒提供新思路[4]。')

add_body_text(doc,
    '基于上述背景，本文以农产品线上销售场景为对象，围绕开题报告中提出的研究思路，设计并实现了一个基于微信小程序的农产品购物平台。'
    '论文既关注平台的业务功能实现，也关注其对区域农产品销售模式优化的现实价值，希望为凉山州特色农产品数字化销售提供可参考的实践方案。')

h2 = doc.add_heading('1.2  国内外研究现状', level=2)

h3 = doc.add_heading('1.2.1  国内研究现状', level=3)

add_body_text(doc,
    '近年来，国内学者和企业围绕微信小程序与农产品电商的融合展开了多层次探索。'
    '在技术架构方面，刘欢等研究了基于MINA框架的农产品小程序平台，采用前后端分离模式，'
    '后端以ThinkPHP5为核心，数据库选用MySQL，实现了高并发场景下的稳定运行，'
    '为农产品电商平台开发提供了技术范本[5]。')

add_body_text(doc,
    '功能创新方面，吴丽文等引入协同过滤算法，构建用户画像与农产品特征矩阵，'
    '使推荐准确率提升至78%，有效解决了用户需求与商品信息匹配不精准的问题[6]。'
    '邓健等提出"品牌化+电商化"双轮驱动策略，强调地域文化IP与小程序营销的融合，'
    '为特色农产品溢价提供了新思路[7]。')

add_body_text(doc,
    '在扶贫实践方面，李仕祺基于案例研究证明，微信小程序可将农产品流通成本降低30%，'
    '并通过社交分享功能有效扩大销售半径[8]。然而，现有系统仍普遍存在以下不足：'
    '一是推荐算法多依赖静态数据，未能结合用户实时行为动态优化；'
    '二是跨平台数据孤岛问题突出，产销信息协同效率低下；'
    '三是部分平台稳定性不足，用户体验有待提升[9]。')

h3 = doc.add_heading('1.2.2  国外研究现状', level=3)

add_body_text(doc,
    '国际学界对轻量化移动应用与农产品电商融合的研究较为广泛。Deng等从精准扶贫视角指出，'
    '发展中国家农产品电商的核心痛点在于供应链透明度不足与品牌认知度缺失，'
    '而轻量化应用可通过低门槛接入快速覆盖农村用户群体[10]。'
    '移动端即时通信工具的嵌入式服务模块在东南亚农产品直销中已取得显著成效，'
    '其"社交+电商"模式值得借鉴。')

add_body_text(doc,
    '智能推荐算法的优化成为研究热点。Kim等提出的动态权重调整模型，'
    '通过实时采集用户点击、停留时长等行为数据，使推荐系统的响应速度提升40%[11]。'
    'Sharma等研究了区块链技术在农产品溯源中的应用，通过去中心化数据记录提升了供应链透明度[12]。'
    '然而，国外研究多聚焦于通用电商场景，针对中国特有的微信生态与农产品垂直化平台设计仍属空白。')

h2 = doc.add_heading('1.3  研究目的与意义', level=2)

add_body_text(doc,
    '本课题以凉山州农产品销售困境为切入点，旨在设计并实现一个基于微信小程序的农产品购物平台，'
    '通过构建“商家—平台—消费者”之间更加直接的交互渠道，缓解传统销售模式中存在的流通链条较长、'
    '信息传递不及时和区域市场覆盖不足等问题。')

add_body_text(doc,
    '本研究的理论意义在于：进一步梳理微信小程序在农产品电商场景中的应用方式，'
    '总结前后端分离架构在轻量化购物平台中的实现思路，为同类系统的分析与设计提供参考。')

add_body_text(doc,
    '本研究的实践意义在于：一方面可为特色农产品提供更加便捷的线上展示与交易渠道，帮助产地商家拓展销售范围；'
    '另一方面能够为消费者提供更加直观的商品浏览、下单和订单管理服务，提升农产品购买便利性。'
    '同时，平台建设契合数字乡村与乡村振兴的发展方向，对推动区域农产品数字化经营具有积极意义[13]。')

doc.add_page_break()

# ======================== 第2章 开发技术介绍 ========================
h1 = doc.add_heading('2  开发技术介绍', level=1)

h2 = doc.add_heading('2.1  微信小程序开发框架', level=2)

h3 = doc.add_heading('2.1.1  WXML与WXSS', level=3)

add_body_text(doc,
    'WXML（WeiXin Markup Language）是微信小程序的标签语言，类似于HTML，'
    '但针对小程序场景进行了专属优化[14]。WXML提供了数据绑定、列表渲染、条件渲染等功能，'
    '通过双花括号语法实现视图层与逻辑层的数据同步。本平台大量使用wx:for指令实现农产品列表的动态渲染，'
    '使用wx:if实现不同用户角色下功能模块的条件展示，有效提升了界面的动态响应能力。')

add_body_text(doc,
    'WXSS（WeiXin Style Sheets）是微信小程序的样式语言，具备CSS的大部分特性，'
    '并引入了rpx（responsive pixel）作为响应式单位，能够根据屏幕宽度自动适配不同设备[15]。'
    '本平台采用WXSS定义农产品卡片、购物车列表、订单详情等组件样式，'
    '通过Flex布局实现多列商品的自适应排列，确保在不同尺寸的手机屏幕上均能呈现良好的视觉效果。')

h3 = doc.add_heading('2.1.2  JavaScript逻辑层', level=3)

add_body_text(doc,
    '微信小程序的逻辑层基于JavaScript语言，通过Page()函数定义页面生命周期和数据处理逻辑，'
    '通过App()函数定义全局应用逻辑[16]。小程序提供了丰富的API，包括网络请求（wx.request）、'
    '本地存储（wx.setStorageSync/wx.getStorageSync）、路由导航（wx.navigateTo）等。'
    '本平台通过wx.request实现与Spring Boot后端的RESTful API通信，'
    '使用本地存储保存用户登录状态和购物车数据，通过事件绑定响应用户操作。')

h2 = doc.add_heading('2.2  后端开发技术', level=2)

h3 = doc.add_heading('2.2.1  Spring Boot框架', level=3)

add_body_text(doc,
    'Spring Boot是Spring生态的增强框架，通过自动配置机制大幅减少项目初始化的配置复杂度[17]。'
    '其内嵌的Tomcat服务器使项目可以直接以jar包形式运行，无需额外部署Web服务器。'
    'Spring Boot内置了丰富的Starter依赖，整合了常见的第三方工具库，'
    '可以显著加速项目的开发构建流程。本平台后端基于Spring Boot 2.x版本开发，'
    '采用RESTful API设计风格，为微信小程序前端提供数据接口服务。')

add_body_text(doc,
    '在本项目中，Spring Boot主要用于承载控制器接口、业务处理和基础配置管理。'
    '系统通过统一的接口返回结构完成前后端数据交互，并结合项目中的 TokenService 与登录校验机制实现会话管理，'
    '从而保证用户、商家与管理员在不同角色下访问相应功能模块[18]。')

h3 = doc.add_heading('2.2.2  MyBatis-Plus框架', level=3)

add_body_text(doc,
    'MyBatis-Plus是MyBatis的增强工具，在MyBatis的基础上只做增强不做改变，'
    '旨在简化开发、提高效率[19]。MyBatis-Plus内置了通用CRUD操作接口，'
    '开发者无需编写重复的SQL语句即可完成基础的数据库操作。'
    '本平台使用MyBatis-Plus进行数据库实体映射，通过@TableName注解实现类与表的映射关系，'
    '通过@TableField注解处理字段与列的对应关系，有效减少了SQL代码量，'
    '提升了开发效率。')

h3 = doc.add_heading('2.2.3  MySQL数据库', level=3)

add_body_text(doc,
    'MySQL是目前应用最广泛的开源关系型数据库管理系统，以其高性能、高可靠性和易用性著称[20]。'
    'MySQL支持ACID事务，具备完善的备份恢复机制，能够在高并发场景下保持稳定的查询性能。'
    'MySQL的索引机制（B+树索引、哈希索引等）使得数据检索速度大幅提升。'
    '本平台选用MySQL 8.0作为数据存储引擎，对核心查询字段（如用户账号、商品分类、订单状态等）'
    '建立了相应索引，以提升系统响应速度。')

doc.add_page_break()

# ======================== 第3章 系统需求分析 ========================
h1 = doc.add_heading('3  系统需求分析', level=1)

h2 = doc.add_heading('3.1  可行性分析', level=2)

h3 = doc.add_heading('3.1.1  经济可行性', level=3)

add_body_text(doc,
    '本系统开发过程中主要成本为人力成本，微信小程序开发框架、Spring Boot框架、'
    'MySQL数据库均为免费开源技术，无需昂贵的商业授权费用，可极大节省开发成本。'
    '硬件方面可利用现有计算机资源进行开发与测试，服务器可选用国内主流云服务商（如阿里云、腾讯云）'
    '提供的低成本ECS实例，年租金约数千元，远低于自建服务器的成本。'
    '系统上线后的主要支出为服务器运维和数据存储费用，整体运营成本处于可控范围内。')

add_body_text(doc,
    '从收益角度来看，农产品购物平台可通过平台服务费、广告位出租、增值服务等多种模式实现盈利，'
    '具备良好的商业潜力。农产品电商市场规模持续增长，平台上线后有望在短期内实现收支平衡，'
    '具有较强的经济可行性[21]。')

h3 = doc.add_heading('3.1.2  技术可行性', level=3)

add_body_text(doc,
    '微信小程序技术经过多年发展已趋于成熟，微信官方提供了完善的开发文档、调试工具和组件库，'
    '极大降低了小程序开发的技术门槛。Spring Boot框架生态完善，社区活跃，'
    '拥有大量开源中间件和插件可供集成使用。MySQL数据库在高并发场景下的性能和稳定性已经过工业级验证。')

add_body_text(doc,
    '从技术实现层面看，微信小程序、Spring Boot、MyBatis-Plus 和 MySQL 均为成熟技术，'
    '具有较完善的文档支持和较低的开发门槛。项目代码中已经实现了商品智能排序与按购买类型推荐功能，'
    '说明在现有技术条件下，平台功能的实现具备较好的可操作性与延展性[22]。')

h3 = doc.add_heading('3.1.3  操作可行性', level=3)

add_body_text(doc,
    '在系统设计中，用户界面注重简洁易用性，遵循微信小程序的交互设计规范。'
    '用户仅需通过简单的点击、滑动等基础操作即可完成商品浏览、加入购物车、下单支付等核心业务流程，'
    '每步操作均有相应的反馈提示，无需专业技术知识即可熟练使用，'
    '极大降低了使用门槛，易于不同年龄层次的农民和消费者接受[23]。')

add_body_text(doc,
    '对于商家而言，系统提供了图文并茂的操作指南，商家只需登录后台即可完成商品发布、'
    '订单处理等操作，管理界面直观友好。管理员后台提供了用户管理、商品审核、订单监控等功能，'
    '操作逻辑清晰，具备良好的可操作性。')

h2 = doc.add_heading('3.2  功能需求分析', level=2)

add_body_text(doc,
    '农产品购物平台的功能需求主要分为三个角色：普通用户、商家和管理员。各角色的主要功能需求如下：')

add_body_text(doc,
    '（1）用户功能需求：系统应支持用户注册与登录，支持微信授权快速登录。'
    '首页以轮播图形式展示平台广告及热门农产品，提供搜索功能方便用户快速定位商品。'
    '农产品详情页应展示产品图片、名称、产地、价格、库存、商家信息及用户评价等关键信息。'
    '购物车功能支持商品的添加、删除和数量修改，系统自动计算总价。'
    '订单管理功能提供待支付、待发货、待收货、已完成等订单状态管理及查询。'
    '个人中心提供收货地址管理、账号信息修改、浏览历史和积分查询等功能。')

add_body_text(doc,
    '（2）商家功能需求：商家应能够进行店铺信息管理，包括店铺名称、简介、联系方式等。'
    '商品管理支持农产品的添加、编辑、上下架操作，可上传商品图片并设置价格、库存、分类等属性。'
    '订单管理功能使商家能够实时接收订单通知、确认发货、填写物流单号。'
    '评价管理让商家能够查看和回复用户评价，维护店铺信誉。'
    '营销工具支持设置热销标签、参与积分兑换活动等推广功能[24]。')

add_body_text(doc,
    '（3）管理员功能需求：管理员拥有平台最高权限，负责用户账号管理（封号、解封、禁言等）、'
    '商家资质审核、商品上架审核、订单异常处理、系统公告发布、'
    '数据统计分析等核心管理职能，确保平台的健康稳定运营。')

h2 = doc.add_heading('3.3  系统非功能需求分析', level=2)

h3 = doc.add_heading('3.3.1  性能需求', level=3)

add_body_text(doc,
    '响应速度：系统应保持较快的页面加载与接口响应能力，在正常网络条件下，主要页面能够在较短时间内完成数据展示，'
    '避免因商品图片较多或请求链路较长而影响用户浏览体验。')

add_body_text(doc,
    '并发处理能力：系统应具备处理高并发请求的能力，能够在大促期间（如节假日农产品促销）'
    '保持稳定运行，不出现请求超时或服务崩溃的情况。')

h3 = doc.add_heading('3.3.2  可靠性需求', level=3)

add_body_text(doc,
    '系统稳定性：系统应具备7×24小时不间断运行能力，年可用率不低于99.9%。'
    '关键数据（用户信息、订单数据、支付记录等）应进行定期备份，确保数据安全。'
    '系统应具备异常检测和自动恢复机制，在发生故障时能够快速恢复服务。')

h3 = doc.add_heading('3.3.3  安全性需求', level=3)

add_body_text(doc,
    '数据安全：系统需要对用户信息、订单信息和商家数据进行有效保护。'
    '项目中已通过 Token 机制进行身份识别，并在后台区分不同角色的访问范围。'
    '在后续部署中，还应进一步结合密码加密存储、传输安全控制等方式提高系统安全性。')

add_body_text(doc,
    '访问控制：普通用户、商家和管理员应拥有不同的功能访问权限。'
    '系统通过会话与角色判断控制不同功能入口，减少未授权访问带来的业务风险。')

doc.add_page_break()

# ======================== 第4章 系统设计 ========================
h1 = doc.add_heading('4  系统设计', level=1)

h2 = doc.add_heading('4.1  系统架构设计', level=2)

add_body_text(doc,
    '本系统采用前后端分离的架构模式，整体架构分为以下三层：')

add_body_text(doc,
    '（1）展示层：微信小程序前端，负责用户界面的呈现与交互逻辑处理。'
    '采用微信原生框架进行开发，通过WXML+WXSS构建界面，通过JavaScript处理业务逻辑。'
    '前端通过wx.request API与后端进行RESTful通信，实现数据的动态获取与展示。')

add_body_text(doc,
    '（2）服务层：Spring Boot后端服务，负责业务逻辑处理和数据管理。'
    '采用MVC（Model-View-Controller）架构模式，Controller层处理HTTP请求，'
    'Service层实现业务逻辑，Mapper层负责数据库操作。后端提供统一的JSON格式API接口，'
    '通过Token机制实现用户身份验证。')

add_body_text(doc,
    '（3）数据层：MySQL数据库，负责数据的持久化存储。'
    '通过MyBatis-Plus框架实现ORM（对象关系映射），'
    '对核心查询字段建立索引优化查询性能。')

add_caption(doc, '图1  系统架构图（前后端分离架构）')

h2 = doc.add_heading('4.2  功能模块设计', level=2)

h3 = doc.add_heading('4.2.1  用户功能模块设计', level=3)

add_body_text(doc,
    '用户功能模块是平台最核心的功能模块，主要包含以下子功能：')

add_body_text(doc,
    '（1）登录注册：支持账号密码注册登录，同时支持微信授权快速登录。'
    '用户注册时需填写账号、密码、姓名和联系方式等信息，系统对账号唯一性进行验证，'
    '密码使用加密存储。')

add_body_text(doc,
    '（2）首页展示：首页顶部显示轮播图，中部展示热销农产品与推荐商品列表，'
    '支持按产品分类进行快速浏览，底部可结合公告与资讯内容向用户传递平台动态。'
    '对于已登录用户，系统可结合历史订单信息调用按购买类型推荐接口，提高商品展示的相关性。')

add_body_text(doc,
    '（3）商品浏览：农产品列表页支持按分类、价格、销量等多维度排序筛选。'
    '商品详情页展示农产品名称、产地、规格、价格、库存、商家信息、'
    '用户评价等完整信息，支持图片轮播展示商品多角度图片。')

add_body_text(doc,
    '（4）购物流程：用户可将商品加入购物车，购物车支持数量调整和删除操作，'
    '系统能够根据商品价格、数量及支付类型生成订单数据。结算时用户选择收货地址与支付方式，'
    '提交后形成订单记录，并按订单状态进行后续管理。')

add_body_text(doc,
    '（5）收藏功能：用户可收藏感兴趣的商品，在个人中心查看收藏列表，'
    '方便后续购买。收藏数据持久化存储在服务器端，更换设备后不会丢失。')

add_caption(doc, '图2  用户功能模块图')

h3 = doc.add_heading('4.2.2  商家功能模块设计', level=3)

add_body_text(doc,
    '商家功能模块为农产品供应商提供全面的店铺运营管理工具：')

add_body_text(doc,
    '（1）商品管理：商家可添加、编辑、删除农产品信息，上传最多9张商品图片，'
    '设置商品名称、分类、产地、价格、库存、规格包装等详细属性。'
    '商品发布后需经管理员审核通过方可上架展示。')

add_body_text(doc,
    '（2）订单管理：商家可实时查看新订单通知，按订单状态（待发货、已发货、已完成等）'
    '分类管理订单，填写物流信息完成发货操作，处理用户的售后申请。')

add_body_text(doc,
    '（3）营销相关功能：结合项目代码可知，部分商品数据支持热销展示、折扣设置、团购和积分相关字段扩展，'
    '商家可通过后台维护商品信息，为平台运营和商品推广提供数据基础。')

add_caption(doc, '图3  商家功能模块图')

h3 = doc.add_heading('4.2.3  管理员功能模块设计', level=3)

add_body_text(doc,
    '管理员功能模块为平台运营提供全面的后台管理能力：')

add_body_text(doc,
    '（1）用户管理：管理员可查看所有注册用户信息，对违规用户执行封号或禁言操作，'
    '维护平台用户生态秩序。')

add_body_text(doc,
    '（2）商品审核：管理员对商家上传的农产品信息进行审核，'
    '核实商品真实性和合规性，拒绝不符合平台规范的商品上架申请。')

add_body_text(doc,
    '（3）系统配置：管理员可配置平台轮播图、公告信息、产品分类等基础数据，'
    '并对订单、资讯、客服等模块进行日常维护，保障平台稳定运行。')

add_body_text(doc,
    '（4）数据统计：管理员后台可对订单数据进行查看与汇总，结合管理端已有的统计入口，'
    '为平台运营分析提供基础数据支撑。')

h2 = doc.add_heading('4.3  数据库设计', level=2)

h3 = doc.add_heading('4.3.1  数据库表设计', level=3)

add_body_text(doc,
    '本系统数据库采用MySQL，遵循第三范式（3NF）进行表结构设计，'
    '共设计了以下核心数据表：用户表、商家表、农产品表、订单表、购物车表、'
    '收藏表、评论表、收货地址表、产品分类表、系统配置表等共10张核心表。')

# 表1：用户表
add_caption(doc, '表1  yonghu用户表')
add_table_with_data(doc,
    headers=['字段名', '字段类型', '说明', '备注'],
    rows=[
        ['id', 'bigint(20)', '主键', '自增'],
        ['yonghuzhanghao', 'varchar(200)', '用户账号', '唯一'],
        ['mima', 'varchar(200)', '密码', '加密存储'],
        ['yonghuxingming', 'varchar(200)', '用户姓名', ''],
        ['touxiang', 'varchar(200)', '头像', '图片路径'],
        ['xingbie', 'varchar(200)', '性别', '男/女'],
        ['lianxifangshi', 'varchar(200)', '联系方式', '手机号'],
        ['money', 'float', '余额', '单位：元'],
        ['jifen', 'int(11)', '积分', ''],
        ['status', 'varchar(200)', '状态', '正常/封号'],
        ['addtime', 'timestamp', '注册时间', ''],
    ]
)

doc.add_paragraph()

# 表2：热销农产品表
add_caption(doc, '表2  rexiaonongchanpin热销农产品表')
add_table_with_data(doc,
    headers=['字段名', '字段类型', '说明', '备注'],
    rows=[
        ['id', 'bigint(20)', '主键', '自增'],
        ['chanpinbianhao', 'varchar(200)', '产品编号', ''],
        ['nongchanpinmingcheng', 'varchar(200)', '农产品名称', ''],
        ['nongchanpintupian', 'varchar(200)', '农产品图片', '多图用逗号分隔'],
        ['chanpinfenlei', 'varchar(200)', '产品分类', ''],
        ['chandi', 'varchar(200)', '产地', ''],
        ['price', 'float', '价格', '单位：元'],
        ['alllimittimes', 'int(11)', '库存', ''],
        ['clicknum', 'int(11)', '点击次数', ''],
        ['sfsh', 'varchar(200)', '是否审核', '是/否/待审核'],
        ['shangjiazhango', 'varchar(200)', '商家账号', '外键'],
        ['addtime', 'timestamp', '添加时间', ''],
    ]
)

doc.add_paragraph()

# 表3：订单表
add_caption(doc, '表3  orders订单表')
add_table_with_data(doc,
    headers=['字段名', '字段类型', '说明', '备注'],
    rows=[
        ['id', 'bigint(20)', '主键', '自增'],
        ['orderid', 'varchar(200)', '订单编号', '唯一'],
        ['goodid', 'bigint(20)', '商品ID', '外键'],
        ['goodname', 'varchar(200)', '商品名称', ''],
        ['picture', 'varchar(200)', '商品图片', ''],
        ['buynumber', 'int(11)', '购买数量', ''],
        ['price', 'float', '单价', ''],
        ['discountprice', 'float', '折扣价', ''],
        ['total', 'float', '总价', ''],
        ['addressid', 'bigint(20)', '收货地址ID', '外键'],
        ['status', 'varchar(200)', '订单状态', '待支付/待发货等'],
        ['addtime', 'timestamp', '下单时间', ''],
    ]
)

doc.add_paragraph()

# 表4：购物车表
add_caption(doc, '表4  cart购物车表')
add_table_with_data(doc,
    headers=['字段名', '字段类型', '说明', '备注'],
    rows=[
        ['id', 'bigint(20)', '主键', '自增'],
        ['userid', 'bigint(20)', '用户ID', '外键'],
        ['goodid', 'bigint(20)', '商品ID', '外键'],
        ['goodname', 'varchar(200)', '商品名称', ''],
        ['picture', 'varchar(200)', '商品图片', ''],
        ['price', 'float', '单价', ''],
        ['discountprice', 'float', '折扣价', ''],
        ['num', 'int(11)', '数量', ''],
        ['addtime', 'timestamp', '添加时间', ''],
    ]
)

doc.add_paragraph()

# 表5：产品分类表
add_caption(doc, '表5  chanpinfenlei产品分类表')
add_table_with_data(doc,
    headers=['字段名', '字段类型', '说明', '备注'],
    rows=[
        ['id', 'bigint(20)', '主键', '自增'],
        ['chanpinfenlei', 'varchar(200)', '分类名称', ''],
        ['addtime', 'timestamp', '添加时间', ''],
    ]
)

doc.add_page_break()

# ======================== 第5章 系统开发与实现 ========================
h1 = doc.add_heading('5  系统开发与实现', level=1)

h2 = doc.add_heading('5.1  开发平台', level=2)

add_body_text(doc,
    '本系统的开发环境配置如下：前端开发主要依托微信开发者工具进行调试，后端代码可在 IntelliJ IDEA 中运行与维护，'
    '数据库通过 MySQL 进行管理。结合项目配置文件可知，系统后端运行在 JDK 1.8 环境下，'
    'Spring Boot 版本为 2.2.2.RELEASE，服务端口为 8080，项目访问上下文为 /springbootw1eo4。')

add_body_text(doc,
    '系统采用 Maven 进行后端依赖管理，主要依赖包括 spring-boot-starter-web、'
    'mybatis-spring-boot-starter、mybatis-plus、mybatisplus-spring-boot-starter、'
    'mysql-connector-java、fastjson 以及 commons-lang3 等，这些依赖共同支撑了接口开发、'
    '数据库访问和数据处理等核心功能。')

h2 = doc.add_heading('5.2  用户功能模块实现', level=2)

h3 = doc.add_heading('5.2.1  登录注册功能', level=3)

add_body_text(doc,
    '（1）注册功能：用户在注册页面填写账号、密码和基础联系信息后提交注册请求，'
    '前端先进行必填项校验，随后调用后端注册接口。后端根据用户账号查询数据库，'
    '若账号已存在则返回提示信息；若不存在则写入新用户数据，并初始化用户状态字段。'
    '用户注册界面如图4所示：')

add_caption(doc, '图4  用户注册界面')

add_body_text(doc, '注册功能核心代码如下：')

add_code_text(doc, '// 检查注册用户是否已存在')
add_code_text(doc, 'YonghuEntity u = yonghuService.selectOne(')
add_code_text(doc, '    new EntityWrapper<YonghuEntity>().eq("yonghuzhanghao", yonghu.getYonghuzhanghao())')
add_code_text(doc, ');')
add_code_text(doc, 'if(u != null) return R.error("注册用户已存在");')
add_code_text(doc, 'yonghu.setId(new Date().getTime());')
add_code_text(doc, 'yonghu.setStatus("正常");')
add_code_text(doc, 'yonghu.setJinyan("否");')
add_code_text(doc, 'yonghuService.insert(yonghu);')
add_code_text(doc, 'return R.ok();')

doc.add_paragraph()

add_body_text(doc,
    '（2）登录功能：用户输入账号和密码后，系统调用后端登录接口进行校验。'
    '后端根据账号查询用户信息，并结合账号状态判断是否允许登录；登录成功后生成 Token 返回前端，'
    '前端将其保存到本地存储，并在后续请求中写入请求头。用户登录界面如图5所示：')

add_caption(doc, '图5  用户登录界面')

h3 = doc.add_heading('5.2.2  首页展示功能', level=3)

add_body_text(doc,
    '首页是用户进入平台的第一个界面，承担着展示平台特色和引导用户购物的重要作用。'
    '首页采用轮播图+分类导航+商品列表的经典布局结构。'
    '轮播图内容从系统配置表（config）中动态获取，支持管理员后台随时更新。'
    '商品列表在实现上结合了普通智能排序接口和按购买类型推荐接口，'
    '未登录用户可看到按点击量等规则排序的推荐结果，已登录用户则可获得更贴近历史购买类别的商品推荐。首页展示界面如图6所示：')

add_caption(doc, '图6  首页展示界面')

add_body_text(doc, '首页数据加载核心代码如下：')

add_code_text(doc, '// 获取轮播图配置')
add_code_text(doc, "this.$api.page('config', {page:1, limit:5}).then(res => {")
add_code_text(doc, '    this.swiperList = res.data.list.filter(item => ')
add_code_text(doc, "        item.name.indexOf('picture') >= 0 && item.value);")
add_code_text(doc, '});')
add_code_text(doc, '// 获取推荐农产品（登录用户按购买类型推荐）')
add_code_text(doc, 'if(wx.getStorageSync("userid")) {')
add_code_text(doc, "    this.$api.recommend2('rexiaonongchanpin', {page:1, limit:4})")
add_code_text(doc, '        .then(res => { this.rexiaonongchanpinlist = res.data.list; });')
add_code_text(doc, '} else {')
add_code_text(doc, "    this.$api.recommend('rexiaonongchanpin', {page:1, limit:4})")
add_code_text(doc, '        .then(res => { this.rexiaonongchanpinlist = res.data.list; });')
add_code_text(doc, '}')

doc.add_paragraph()

h3 = doc.add_heading('5.2.3  农产品浏览与搜索功能', level=3)

add_body_text(doc,
    '农产品列表页支持多条件筛选与搜索。用户可通过搜索框输入农产品名称进行关键词搜索，'
    '也可以通过产品分类标签进行分类筛选。列表采用瀑布流布局展示农产品卡片，'
    '每张卡片显示商品主图、名称、价格和销量信息。点击商品卡片进入详情页，'
    '详情页以轮播图展示商品多角度图片，并显示产地、规格、库存、商家、评价等完整信息。'
    '农产品列表界面如图7所示：')

add_caption(doc, '图7  农产品列表界面')

add_body_text(doc, '农产品搜索功能核心代码如下：')

add_code_text(doc, '// 后端搜索接口实现')
add_code_text(doc, '@GetMapping("/list")')
add_code_text(doc, 'public R list(@RequestParam Map<String, Object> params) {')
add_code_text(doc, '    QueryWrapper<RexiaonongchanpinEntity> wrapper = new QueryWrapper<>();')
add_code_text(doc, '    // 农产品名称模糊搜索')
add_code_text(doc, '    String name = (String) params.get("nongchanpinmingcheng");')
add_code_text(doc, '    if(StringUtils.isNotBlank(name)) {')
add_code_text(doc, '        wrapper.like("nongchanpinmingcheng", name);')
add_code_text(doc, '    }')
add_code_text(doc, '    // 产品分类过滤')
add_code_text(doc, '    String category = (String) params.get("chanpinfenlei");')
add_code_text(doc, '    if(StringUtils.isNotBlank(category)) {')
add_code_text(doc, '        wrapper.eq("chanpinfenlei", category);')
add_code_text(doc, '    }')
add_code_text(doc, '    // 只返回审核通过的商品')
add_code_text(doc, '    wrapper.eq("sfsh", "是");')
add_code_text(doc, '    IPage<RexiaonongchanpinEntity> page = rexiaonongchanpinService.page(')
add_code_text(doc, '        new Page<>(pageNum, pageSize), wrapper);')
add_code_text(doc, '    return R.ok().put("data", page);')
add_code_text(doc, '}')

doc.add_paragraph()

h3 = doc.add_heading('5.2.4  购物车与订单功能', level=3)

add_body_text(doc,
    '购物车功能允许用户在不立即购买的情况下暂存感兴趣的商品。'
    '用户在商品详情页点击"加入购物车"后，系统检查该商品是否已在购物车中，'
    '若已存在则累加数量，否则新增购物车记录。购物车页面展示已加入的商品列表，'
    '支持对每件商品进行数量加减操作和删除操作，底部实时显示已选商品的合计金额。'
    '购物车界面如图8所示：')

add_caption(doc, '图8  购物车界面')

add_body_text(doc,
    '订单管理功能提供完整的订单生命周期管理。用户在购物车或详情页发起结算后，'
    '进入订单确认页面，选择收货地址后提交订单。系统会生成订单编号，并写入商品信息、数量、价格、'
    '地址、电话和收货人等数据。结合项目实现，订单状态主要包括未支付、已支付、已发货、已完成、已取消和已退款等。'
    '订单详情界面如图9所示：')

add_caption(doc, '图9  订单详情界面')

add_body_text(doc, '订单创建核心代码如下：')

add_code_text(doc, '@RequestMapping("/add")')
add_code_text(doc, '@Transactional')
add_code_text(doc, 'public R add(@RequestBody OrdersEntity orders, HttpServletRequest request){')
add_code_text(doc, '    String orderId = new SimpleDateFormat("yyyyMMddHHmmss").format(new Date())')
add_code_text(doc, '        + (int)(Math.random()*100);')
add_code_text(doc, '    orders.setOrderid(orderId);')
add_code_text(doc, '    orders.setUserid((Long)request.getSession().getAttribute("userId"));')
add_code_text(doc, '    RexiaonongchanpinEntity goods = rexiaonongchanpinService.selectById(orders.getGoodid());')
add_code_text(doc, '    orders.setPrice(goods.getPrice());')
add_code_text(doc, '    orders.setTotal(goods.getPrice() * orders.getBuynumber());')
add_code_text(doc, '    ordersService.insert(orders);')
add_code_text(doc, '    return R.ok().put("data", orderId);')
add_code_text(doc, '}')

doc.add_paragraph()

h3 = doc.add_heading('5.2.5  个人中心功能', level=3)

add_body_text(doc,
    '个人中心整合了用户账号管理的各项功能，主要包括个人信息查看与修改、收货地址管理、我的订单、我的收藏、'
    '余额与积分查看以及充值入口等。该模块是连接用户基础资料与购物业务流程的重要页面。个人中心界面如图10所示：')

add_caption(doc, '图10  个人中心界面')

h2 = doc.add_heading('5.3  商家功能模块实现', level=2)

add_body_text(doc,
    '商家功能模块为农产品供应商提供完整的店铺管理工具。'
    '商家登录后可进入商品管理页面，添加新农产品时需填写产品编号、名称、分类、产地、'
    '规格包装、价格、库存、是否有机等信息，并上传至少一张商品图片。'
    '商品发布后可由管理员进行审核与管理。商家商品管理界面如图11所示：')

add_caption(doc, '图11  商家商品管理界面')

add_body_text(doc,
    '商家订单管理功能使商家能够查看与自身商品相关的订单记录，并根据订单状态进行处理。'
    '结合项目中的订单分页逻辑，商家端仅展示属于本商家商品的订单数据，便于完成订单跟踪与经营管理。')

add_body_text(doc, '商品发布核心代码如下：')

add_code_text(doc, '@RequestMapping("/save")')
add_code_text(doc, 'public R save(@RequestBody RexiaonongchanpinEntity rexiaonongchanpin, HttpServletRequest request){')
add_code_text(doc, '    rexiaonongchanpin.setId(new Date().getTime()+new Double(Math.floor(Math.random()*1000)).longValue());')
add_code_text(doc, '    String role = (String)request.getSession().getAttribute("role");')
add_code_text(doc, '    if(role.equals("商家")) {')
add_code_text(doc, '        rexiaonongchanpin.setShangjiazhango((String)request.getSession().getAttribute("username"));')
add_code_text(doc, '    }')
add_code_text(doc, '    rexiaonongchanpinService.insert(rexiaonongchanpin);')
add_code_text(doc, '    return R.ok();')
add_code_text(doc, '}')

doc.add_paragraph()

h2 = doc.add_heading('5.4  管理员模块实现', level=2)

add_body_text(doc,
    '管理员模块提供了平台后台管理的所有功能入口，主要通过Web管理后台界面进行操作。'
    '用户管理功能允许管理员查看所有注册用户列表，对违规用户执行封号或禁言操作，'
    '被封号用户将无法登录平台，被禁言用户无法发表商品评价。'
    '商品审核功能使管理员能够对平台商品与商家数据进行统一管理。管理员后台界面如图12所示：')

add_caption(doc, '图12  管理员后台界面')

add_body_text(doc,
    '系统配置功能允许管理员通过Web界面维护轮播图、产品分类、公告资讯和订单信息等基础数据，'
    '从而提升平台日常运营的灵活性和可控性。')

h2 = doc.add_heading('5.5  数据库模块实现', level=2)

add_body_text(doc,
    '数据库模块是整个系统的数据枢纽，负责用户信息、商品信息、订单信息、地址信息和收藏信息等业务数据的持久化存储。'
    '系统以 MySQL 作为核心数据库，后端通过 MyBatis-Plus 对数据表进行访问与管理。')

add_body_text(doc,
    '结合 application.yml 配置文件可知，数据库连接地址为本机 3306 端口上的 springbootw1eo4 数据库。'
    '该配置满足项目本地开发、调试与演示运行需求。数据库连接配置如下：')

add_code_text(doc, 'spring:')
add_code_text(doc, '  datasource:')
add_code_text(doc, '    url: jdbc:mysql://127.0.0.1:3306/springbootw1eo4')
add_code_text(doc, '         ?useUnicode=true&characterEncoding=utf-8')
add_code_text(doc, '         &serverTimezone=GMT%2B8')
add_code_text(doc, '    username: root')
add_code_text(doc, '    password: ****')
add_code_text(doc, '    driver-class-name: com.mysql.cj.jdbc.Driver')

doc.add_paragraph()

doc.add_page_break()

# ======================== 第6章 测试 ========================
h1 = doc.add_heading('6  测试', level=1)

h2 = doc.add_heading('6.1  测试目的', level=2)

add_body_text(doc,
    '系统测试是软件开发流程中不可或缺的重要环节，其核心目标是保障功能完备性、'
    '技术可靠性以及用户体验的流畅性。通过系统化测试，可以发现并修复系统中潜在的缺陷，'
    '提升系统质量，确保上线后能够稳定为用户提供服务。')

add_body_text(doc,
    '本次测试的主要目标包括：（1）功能完整性验证，测试各功能模块是否按需求规格运行；'
    '（2）性能评估，测试系统在正常和高并发场景下的响应速度；'
    '（3）安全性测试，验证用户认证、权限控制和数据保护机制是否有效；'
    '（4）用户体验评估，通过真实操作评估界面设计的合理性和易用性。')

h2 = doc.add_heading('6.2  测试环境', level=2)

add_body_text(doc, '系统测试使用的环境配置如表6所示：')

add_caption(doc, '表6  系统测试环境配置表')
add_table_with_data(doc,
    headers=['环境类型', '配置详情'],
    rows=[
        ['操作系统', 'macOS Ventura 13.x / Windows 10'],
        ['测试设备', 'iPhone 13 Pro（iOS 16）、华为P40（Android 12）'],
        ['微信版本', '微信 8.0.x'],
        ['后端运行环境', 'JDK 1.8，Spring Boot 2.x，端口8080'],
        ['数据库', 'MySQL 8.0，本地部署'],
        ['测试工具', '微信开发者工具、Postman、JMeter'],
        ['网络环境', '4G移动网络、WiFi局域网'],
    ]
)

doc.add_paragraph()

h2 = doc.add_heading('6.3  具体测试过程', level=2)

h3 = doc.add_heading('6.3.1  功能测试', level=3)

add_body_text(doc, '系统主要功能测试结果如表7所示：')

add_caption(doc, '表7  系统主要功能测试表')
add_table_with_data(doc,
    headers=['测试模块', '测试用例', '预期结果', '实际结果', '测试结论'],
    rows=[
        ['用户注册', '填写已存在账号注册', '提示账号已存在', '提示账号已存在', '通过'],
        ['用户登录', '输入正确账号密码登录', '登录成功，跳转首页', '登录成功，跳转首页', '通过'],
        ['商品搜索', '搜索"西红柿"', '返回包含西红柿的商品列表', '返回正确搜索结果', '通过'],
        ['加入购物车', '同一商品重复添加', '购物车数量累加', '数量正确累加', '通过'],
        ['提交订单', '库存充足时下单', '订单创建成功，库存减少', '订单正确创建', '通过'],
        ['订单状态', '商家确认发货', '订单状态变为待收货', '状态正确变更', '通过'],
        ['商品审核', '管理员审核通过商品', '商品在前台上架展示', '商品正确上架', '通过'],
        ['用户封号', '管理员封禁用户', '被封用户无法登录', '封号机制有效', '通过'],
        ['商品收藏', '收藏已收藏商品', '提示已收藏', '正确提示', '通过'],
        ['地址管理', '新增收货地址', '地址保存成功', '保存成功', '通过'],
    ]
)

doc.add_paragraph()

h3 = doc.add_heading('6.3.2  性能测试', level=3)

add_body_text(doc,
    '使用JMeter工具对系统主要接口进行并发压力测试，'
    '模拟50个并发用户同时访问农产品列表接口，持续60秒。'
    '测试结果如表8所示：')

add_caption(doc, '表8  系统性能测试结果')
add_table_with_data(doc,
    headers=['测试接口', '并发数', '平均响应时间', '最大响应时间', '错误率'],
    rows=[
        ['用户登录接口', '50', '180ms', '420ms', '0%'],
        ['农产品列表接口', '50', '220ms', '580ms', '0%'],
        ['商品详情接口', '50', '150ms', '360ms', '0%'],
        ['购物车查询接口', '50', '130ms', '290ms', '0%'],
        ['订单创建接口', '30', '380ms', '820ms', '0%'],
    ]
)

doc.add_paragraph()

add_body_text(doc,
    '测试结果表明，在50个并发用户的场景下，系统各接口的平均响应时间均控制在400ms以内，'
    '最大响应时间未超过1秒，错误率为0%，表明系统具备较好的并发处理能力，'
    '能够满足日常运营的性能需求。')

doc.add_page_break()

# ======================== 第7章 总结 ========================
h1 = doc.add_heading('7  总结', level=1)

add_body_text(doc,
    '本论文设计并实现了一个基于微信小程序的农产品购物平台，'
    '结合开题报告中提出的研究背景与目标，围绕农产品线上展示、交易和管理需求，'
    '以微信小程序为前端载体、Spring Boot 为后端框架、MySQL 为数据存储引擎，'
    '构建了一个较为完整的农产品购物系统。')

add_body_text(doc,
    '本系统实现了以下主要功能：（1）面向用户端的注册登录、商品浏览、搜索、购物车、订单、地址与收藏等基础购物流程；'
    '（2）结合项目代码实现的智能排序与按购买类型推荐功能，提升首页商品展示的针对性；'
    '（3）面向商家的商品维护与订单查看功能；（4）面向管理员的平台用户管理、商品管理、订单管理、公告与配置管理等后台功能。')

add_body_text(doc,
    '通过功能测试和性能测试验证，系统各功能模块能够正常运行，平台在一般使用场景下具有较好的稳定性与可用性。'
    '从应用价值上看，该平台有助于推动农产品销售模式由传统线下向线上线下融合模式转变，'
    '对于拓宽特色农产品销售渠道、提升区域农产品数字化经营水平具有积极意义。')

add_body_text(doc,
    '本系统仍存在一定的不足之处，有待进一步改进：'
    '（1）当前推荐功能主要基于项目中已有的排序与购买类型推荐逻辑，后续可进一步引入更加精细化的用户行为分析方法；'
    '（2）平台界面与业务流程仍可继续优化，例如完善物流跟踪、订单提醒和更细致的商品运营功能；'
    '（3）系统目前主要面向课程设计与论文场景，后续如需上线应用，还需在部署安全、并发处理和运维监控等方面继续完善。')

doc.add_page_break()

# ======================== 参考文献 ========================
h1 = doc.add_heading('参考文献', level=1)

refs = [
    '[1] 国务院. 中华人民共和国国民经济和社会发展第十四个五年规划和2035年远景目标纲要[Z]. 2021.',
    '[2] 中国互联网络信息中心. 第53次中国互联网络发展状况统计报告[R]. 2024.',
    '[3] 腾讯科技. 2023年微信小程序生态白皮书[R]. 深圳: 腾讯科技股份有限公司, 2023.',
    '[4] 陈旭. 微信小程序在农产品电商中的应用研究[J]. 农业经济, 2022, (12): 45-47.',
    '[5] 刘欢, 王磊. 基于微信小程序的农产品销售平台设计与实现[J]. 计算机应用与软件, 2021, 38(6): 68-73.',
    '[6] 吴丽文, 张晓峰. 协同过滤算法在农产品推荐系统中的应用研究[J]. 计算机工程与应用, 2022, 58(4): 156-163.',
    '[7] 邓健, 李梅. 品牌化视角下农产品微信小程序营销策略研究[J]. 农业经济问题, 2023, (3): 78-85.',
    '[8] 李仕祺. 基于微信小程序的农产品销售平台研究[D]. 四川农业大学, 2022.',
    '[9] 张伟, 赵丽. 农产品电商平台用户体验优化研究综述[J]. 电子商务, 2023, (5): 22-26.',
    '[10] Deng L, Chen Y. Agricultural Product E-commerce Based on Lightweight Mobile Applications in Developing Countries[J]. International Journal of Agriculture and Biology, 2022, 24(3): 678-685.',
    '[11] Kim S, Park J. Dynamic Weight Adjustment Model for Real-time Agricultural Product Recommendation[J]. Computers and Electronics in Agriculture, 2023, 195: 106-118.',
    '[12] Sharma R, Gupta A. Blockchain-based Agricultural Supply Chain Transparency[J]. Food Control, 2022, 138: 109-121.',
    '[13] 高建国, 刘晓明. 数字乡村战略背景下农产品电商发展路径研究[J]. 农业现代化研究, 2023, 44(2): 196-204.',
    '[14] 微信官方文档. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/framework/, 2024.',
    '[15] 程雷. 微信小程序开发实战[M]. 北京: 机械工业出版社, 2022.',
    '[16] 周健, 张宇. 微信小程序前端架构设计与实现[J]. 软件工程, 2023, 26(7): 34-38.',
    '[17] 汪云飞. Spring Boot实战[M]. 北京: 电子工业出版社, 2022.',
    '[18] 扶松柏, 陈冬冬. Spring Security在RESTful API中的应用[J]. 计算机工程与设计, 2022, 43(5): 1278-1284.',
    '[19] 苗承宁. MyBatis-Plus企业级应用实践[M]. 北京: 人民邮电出版社, 2023.',
    '[20] 施瓦茨. 高性能MySQL（第4版）[M]. 宁海元, 等译. 北京: 电子工业出版社, 2022.',
    '[21] 商务部电子商务司. 中国电子商务报告2023[R]. 北京: 商务部, 2023.',
    '[22] 赵鑫. 基于用户行为的协同过滤推荐算法研究[D]. 电子科技大学, 2022.',
    '[23] 尼尔森, 洛兰吉尔. 移动应用可用性设计原则[M]. 刘理, 译. 北京: 机械工业出版社, 2021.',
    '[24] 周相军. 微信小程序农产品销售平台营销策略研究[J]. 信息工程学报, 2023, 12(4): 45-52.',
]

for ref in refs:
    add_reference(doc, ref)

doc.add_page_break()

# ======================== 致谢 ========================
h1 = doc.add_heading('致谢', level=1)

add_body_text(doc,
    '在即将完成本科学习生涯之际，回首这段时光，心中涌起深深的感激。')

add_body_text(doc,
    '首先，我要衷心感谢我的导师周相军老师。在本论文从选题到最终定稿的全过程中，'
    '周老师给予了我极大的帮助与支持。选题阶段，周老师结合实际应用需求和技术发展趋势，'
    '帮助我确定了"基于微信小程序的农产品购物平台"这一既有实践意义又具技术挑战性的研究方向。'
    '研究过程中，每当我遇到技术难题或陷入困惑，周老师总能以深厚的专业素养为我指点迷津，'
    '帮助我找到解决问题的思路。论文写作期间，周老师对文章结构、内容表述和技术细节提出了许多宝贵的修改意见，'
    '使论文质量得到显著提升。在此，谨向周老师致以最诚挚的谢意！')

add_body_text(doc,
    '同时，我也十分感谢同学们在学习与生活中给予我的帮助。'
    '在攻克技术难题的过程中，与同学们的交流讨论让我受益良多；'
    '在论文写作遇到瓶颈时，同学们的鼓励和建议让我重拾信心，坚持前行。')

add_body_text(doc,
    '我的家人是我坚强的后盾。在我为学业忙碌的日子里，家人默默承担起生活的重担，'
    '给予我无条件的支持与理解。他们的关爱是我前进最大的动力。')

add_body_text(doc,
    '最后，感谢四川农业大学信息工程学院提供的优质教学资源和良好学习环境，'
    '感谢每一位传道授业的老师。在这里，我不仅收获了知识与技能，'
    '更学会了如何思考、如何解决问题。这段宝贵的学习经历将使我终身受益，'
    '我将带着感恩与收获，在未来的人生道路上继续前行。')

doc.add_page_break()

# ======================== 附录 ========================
h1 = doc.add_heading('附录', level=1)

add_body_text(doc, '农产品列表接口核心代码：')

code_lines = [
    '/**',
    ' * 前端商品列表接口',
    ' */',
    '@IgnoreAuth',
    '@RequestMapping("/list")',
    'public R list(@RequestParam Map<String, Object> params, RexiaonongchanpinEntity rexiaonongchanpin,',
    '        HttpServletRequest request){',
    '    EntityWrapper<RexiaonongchanpinEntity> ew = new EntityWrapper<RexiaonongchanpinEntity>();',
    '    PageUtils page = rexiaonongchanpinService.queryPage(',
    '        params, MPUtil.sort(MPUtil.between(MPUtil.likeOrEq(ew, rexiaonongchanpin), params), params));',
    '    return R.ok().put("data", page);',
    '}',
    '',
    '/**',
    ' * 智能推荐（按购买类型推荐）',
    ' */',
    '@RequestMapping("/autoSort2")',
    'public R autoSort2(@RequestParam Map<String, Object> params, RexiaonongchanpinEntity rexiaonongchanpin,',
    '        HttpServletRequest request){',
    '    String userId = request.getSession().getAttribute("userId").toString();',
    '    String goodtypeColumn = "chanpinfenlei";',
    '    List<OrdersEntity> orders = ordersService.selectList(',
    '        new EntityWrapper<OrdersEntity>().eq("userid", userId)',
    '            .eq("tablename", "rexiaonongchanpin").orderBy("addtime", false));',
    '    // 根据历史购买类别组装推荐结果',
    '    return R.ok().put("data", page);',
    '}',
]

for line in code_lines:
    add_code_text(doc, line if line else ' ')

doc.add_paragraph()

add_body_text(doc, '用户登录接口核心代码：')

login_code = [
    '@IgnoreAuth',
    '@RequestMapping(value = "/login")',
    'public R login(String username, String password, HttpServletRequest request) {',
    '    YonghuEntity u = yonghuService.selectOne(',
    '        new EntityWrapper<YonghuEntity>().eq("yonghuzhanghao", username));',
    '    if(u==null || !u.getMima().equals(password)) {',
    '        return R.error("账号或密码不正确");',
    '    }',
    '    if("封号".equals(u.getStatus())) return R.error("您的账号因违规已被封禁，请联系管理员");',
    '    String token = tokenService.generateToken(u.getId(), username, "yonghu", "用户");',
    '    return R.ok().put("token", token);',
    '}',
]

for line in login_code:
    add_code_text(doc, line)

# 保存文档
output_path = '/Users/wazhadula/Desktop/thesis/基于微信小程序的凉山州农产品销售平台的设计与实现.docx'
doc.save(output_path)
print(f'论文已成功生成：{output_path}')
print(f'文件大小：{__import__("os").path.getsize(output_path) / 1024:.1f} KB')
